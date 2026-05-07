import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BANK_PATH = ROOT / "data" / "question-bank.json"
OUT_PATH = ROOT / "毛概题库_排版复习版.docx"

FONT_BODY = "Microsoft YaHei"
COLOR_BLUE = RGBColor(31, 91, 142)
COLOR_DARK = RGBColor(38, 38, 38)
COLOR_MUTED = RGBColor(95, 108, 120)


def set_run_font(run, name=FONT_BODY, size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def type_label(question):
    raw = str(question.get("type") or "")
    if "Multiple" in raw or "多选" in raw or raw == "2":
        return "多选题"
    if "Judgement" in raw or "判断" in raw or raw == "6":
        return "判断题"
    return "单选题"


def answer_labels(question):
    labels = [str(label).upper() for label in question.get("correctLabels") or []]
    if labels:
        return sorted(set(labels))

    texts = [str(text).strip() for text in question.get("correctTexts") or []]
    resolved = []
    for text in texts:
        compact = text.lower().replace(" ", "")
        for option in question.get("options") or []:
            option_text = str(option.get("text") or "").strip()
            option_key = option_text.lower().replace(" ", "")
            if compact in {option_key, str(option.get("label") or "").lower()}:
                resolved.append(str(option.get("label")).upper())
            elif compact == "true" and ("正确" in option_text or "对" in option_text):
                resolved.append(str(option.get("label")).upper())
            elif compact == "false" and ("错误" in option_text or "错" in option_text):
                resolved.append(str(option.get("label")).upper())
    return sorted(set(resolved))


def answer_text(question):
    labels = answer_labels(question)
    options = {str(option.get("label")).upper(): option for option in question.get("options") or []}
    pieces = []
    for label in labels:
        option = options.get(label)
        if option and option.get("text"):
            pieces.append(f"{label}. {option['text']}")
        else:
            pieces.append(label)
    if pieces:
        return "；".join(pieces)
    return "；".join(str(text) for text in question.get("correctTexts") or [])


def answer_short(question):
    labels = answer_labels(question)
    if labels:
        return "、".join(labels)
    return "、".join(str(text) for text in question.get("correctTexts") or [])


def add_plain_run(paragraph, text, size=10.5, bold=False, color=COLOR_DARK):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def add_heading_band(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    set_paragraph_shading(p, "DDEBF7")
    add_plain_run(p, f"  {title}", size=15, bold=True, color=COLOR_BLUE)
    if subtitle:
        add_plain_run(p, f"  {subtitle}", size=9, color=COLOR_MUTED)


def add_quick_index(doc, questions):
    add_heading_band(doc, "答案速查", "按题库原始编号排列")
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell_index, cell in enumerate(header.cells):
        set_cell_shading(cell, "1F5B8E")
        set_cell_margins(cell, top=90, bottom=90, start=80, end=80)
        text = "题号" if cell_index % 2 == 0 else "答案"
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=9, bold=True, color=RGBColor(255, 255, 255))

    for row_start in range(0, len(questions), 4):
        row = table.add_row()
        for block in range(4):
            q_index = row_start + block
            number_cell = row.cells[block * 2]
            answer_cell = row.cells[block * 2 + 1]
            for cell in (number_cell, answer_cell):
                set_cell_margins(cell, top=70, bottom=70, start=80, end=80)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if q_index >= len(questions):
                number_cell.text = ""
                answer_cell.text = ""
                continue
            number_cell.text = ""
            p_num = number_cell.paragraphs[0]
            p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_plain_run(p_num, f"{q_index + 1:03d}", size=9, bold=True, color=COLOR_BLUE)
            answer_cell.text = ""
            p_ans = answer_cell.paragraphs[0]
            p_ans.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_plain_run(p_ans, answer_short(questions[q_index]), size=9.5, bold=True, color=COLOR_DARK)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, 1.25 if idx % 2 == 0 else 2.0)


def add_question(doc, question, original_index):
    labels = set(answer_labels(question))
    kind = type_label(question)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(9)
    meta.paragraph_format.space_after = Pt(2)
    meta.paragraph_format.keep_with_next = True
    add_plain_run(meta, f"Q{original_index:03d}", size=9, bold=True, color=RGBColor(255, 255, 255))
    set_paragraph_shading(meta, "1F5B8E")
    add_plain_run(meta, f"  {kind}", size=9, bold=True, color=RGBColor(255, 255, 255))

    stem = doc.add_paragraph()
    stem.paragraph_format.space_before = Pt(2)
    stem.paragraph_format.space_after = Pt(4)
    stem.paragraph_format.keep_with_next = True
    stem.paragraph_format.line_spacing = 1.18
    add_plain_run(stem, question.get("stem") or "", size=11, bold=True, color=COLOR_DARK)

    for option in question.get("options") or []:
        label = str(option.get("label") or "").upper()
        opt = doc.add_paragraph()
        opt.paragraph_format.left_indent = Cm(0.45)
        opt.paragraph_format.first_line_indent = Cm(-0.25)
        opt.paragraph_format.space_after = Pt(1.2)
        opt.paragraph_format.line_spacing = 1.12
        is_answer = label in labels
        add_plain_run(opt, f"{label}. ", size=10.2, bold=True, color=COLOR_BLUE if is_answer else COLOR_MUTED)
        add_plain_run(opt, str(option.get("text") or ""), size=10.2, bold=is_answer, color=COLOR_BLUE if is_answer else COLOR_DARK)

    answer = doc.add_paragraph()
    answer.paragraph_format.left_indent = Cm(0.25)
    answer.paragraph_format.space_before = Pt(3)
    answer.paragraph_format.space_after = Pt(5)
    answer.paragraph_format.keep_together = True
    set_paragraph_shading(answer, "EAF3F8")
    add_plain_run(answer, "  答案：", size=10, bold=True, color=COLOR_BLUE)
    add_plain_run(answer, answer_text(question), size=10, bold=True, color=COLOR_DARK)


def build_docx():
    bank = json.loads(BANK_PATH.read_text(encoding="utf-8"))
    questions = list(bank.get("questions", {}).values())

    counts = {"单选题": 0, "多选题": 0, "判断题": 0}
    for question in questions:
        counts[type_label(question)] = counts.get(type_label(question), 0) + 1

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

    styles = doc.styles
    styles["Normal"].font.name = FONT_BODY
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("毛概题库复习版")
    set_run_font(run, size=23, bold=True, color=COLOR_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    add_plain_run(
        subtitle,
        f"共 {len(questions)} 题 | 单选 {counts.get('单选题', 0)} | 多选 {counts.get('多选题', 0)} | 判断 {counts.get('判断题', 0)} | 生成时间 {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        size=10,
        color=COLOR_MUTED,
    )

    add_quick_index(doc, questions)
    doc.add_page_break()

    indexed = list(enumerate(questions, start=1))
    for group_name in ("单选题", "多选题", "判断题"):
        group = [(idx, question) for idx, question in indexed if type_label(question) == group_name]
        if not group:
            continue
        add_heading_band(doc, group_name, f"共 {len(group)} 题")
        for idx, question in group:
            add_question(doc, question, idx)

    doc.core_properties.title = "毛概题库复习版"
    doc.core_properties.subject = "雨课堂题库整理"
    doc.core_properties.author = "Codex"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_docx()
